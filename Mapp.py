import base64
import io
import json
import os
import re
import tempfile
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import pypandoc
import requests
import streamlit as st
import streamlit.components.v1 as components

# --- CẤU HÌNH GIAO DIỆN TRANG ---
st.set_page_config(
    page_title="MinerU JSON Converter Pro", 
    page_icon="🚀", 
    layout="wide"
)

# --- CSS TÙY CHỈNH GIAO DIỆN ---
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .preview-card { 
        background-color: white; 
        padding: 25px; 
        border-radius: 12px; 
        border: 1px solid #e9ecef;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.05);
        margin-top: 20px;
    }
    .stDownloadButton button { width: 100%; border-radius: 8px; font-weight: 600; }
    </style>
""", unsafe_allow_html=True)


# --- 1. CÁC HÀM XỬ LÝ DÙNG CHUNG ---

def get_image_bytes(img_path_str, uploaded_images_map, json_upload_dir=""):
    if not img_path_str:
        return None
    
    if img_path_str.startswith("http://") or img_path_str.startswith("https://"):
        try:
            response = requests.get(img_path_str, timeout=10)
            if response.status_code == 200:
                return io.BytesIO(response.content)
        except Exception:
            pass
        return None

    clean_name = os.path.basename(img_path_str)
    
    if uploaded_images_map and clean_name in uploaded_images_map:
        return io.BytesIO(uploaded_images_map[clean_name])
    
    if json_upload_dir:
        auto_path = os.path.join(json_upload_dir, "images", clean_name)
        if os.path.exists(auto_path):
            with open(auto_path, "rb") as f:
                return io.BytesIO(f.read())
                
    fallback_paths = [
        os.path.join("images", clean_name),
        os.path.join(os.getcwd(), "images", clean_name),
        clean_name
    ]
    for path in fallback_paths:
        if os.path.exists(path):
            with open(path, "rb") as f:
                return io.BytesIO(f.read())
                
    return None

def format_latex_string(latex_str):
    if not latex_str:
        return ""
    latex_clean = latex_str.strip()
    if latex_clean.startswith("$") and latex_clean.endswith("$"):
        latex_clean = latex_clean[1:-1].strip()
    return f"${latex_clean}$"


# --- 2. XỬ LÝ CHUYỂN JSON SANG MARKDOWN ---

def process_html_table_md(table_html, uploaded_images_map, temp_dir, json_upload_dir=""):
    soup = BeautifulSoup(table_html, "html.parser")
    rows = soup.find_all("tr")
    if not rows:
        return ""

    md_table = []
    headers_parsed = False
    img_counter = 0

    for tr in rows:
        cells = tr.find_all(["td", "th"])
        row_content = []
        for cell in cells:
            cell_text = ""
            for node in cell.children:
                if node.name == "eq":
                    cell_text += f" {format_latex_string(node.get_text())} "
                elif node.name == "img":
                    img_src = node.get("src")
                    if img_src and temp_dir:
                        img_stream = get_image_bytes(img_src, uploaded_images_map, json_upload_dir)
                        if img_stream:
                            img_counter += 1
                            img_path = os.path.join(temp_dir, f"tbl_img_{img_counter}.png")
                            with open(img_path, "wb") as f:
                                f.write(img_stream.getvalue())
                            cell_text += f" ![]({img_path}) "
                elif node.name is None:
                    cell_text += str(node).strip()
            row_content.append(cell_text.strip().replace("\n", " "))

        md_row = "| " + " | ".join(row_content) + " |"
        md_table.append(md_row)

        if not headers_parsed:
            separator = "| " + " | ".join(["---"] * len(row_content)) + " |"
            md_table.append(separator)
            headers_parsed = True

    return "\n".join(md_table) + "\n\n"

def convert_json_to_markdown(json_data, uploaded_images_map, temp_dir, json_upload_dir=""):
    md_lines = []
    pdf_info = json_data.get("pdf_info", [])
    img_counter = 0

    for page in pdf_info:
        para_blocks = page.get("para_blocks", [])

        for block in para_blocks:
            b_type = block.get("type")

            if b_type in ["text", "title"]:
                p_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        span_type = span.get("type")
                        content = span.get("content", "")

                        if span_type == "text":
                            if re.match(r"^Bài\s+\d+", content.strip()):
                                p_text += f"**{content}**"
                            else:
                                p_text += content
                        elif span_type == "inline_equation":
                            p_text += f" {format_latex_string(content)} "

                if p_text.strip():
                    md_lines.append(p_text.strip() + "\n\n")

            elif b_type in ["image", "chart", "figure"]:
                paths_to_check = []
                if block.get("image_path"):
                    paths_to_check.append(block.get("image_path"))
                for sub_b in block.get("blocks", []):
                    if sub_b.get("image_path"):
                        paths_to_check.append(sub_b.get("image_path"))
                    for line in sub_b.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("image_path"):
                                paths_to_check.append(span.get("image_path"))

                for img_path_str in paths_to_check:
                    img_stream = get_image_bytes(img_path_str, uploaded_images_map, json_upload_dir)
                    if img_stream and temp_dir:
                        img_counter += 1
                        local_img_path = os.path.join(temp_dir, f"img_{img_counter}.png")
                        with open(local_img_path, "wb") as f:
                            f.write(img_stream.getvalue())
                        md_lines.append(f"![Hình ảnh]({local_img_path})\n\n")

            elif b_type == "table":
                for sub_b in block.get("blocks", []):
                    for line in sub_b.get("lines", []):
                        for span in line.get("spans", []):
                            table_html = span.get("html")
                            if table_html:
                                md_table_str = process_html_table_md(table_html, uploaded_images_map, temp_dir, json_upload_dir)
                                md_lines.append(md_table_str)

    return "".join(md_lines)


# --- 3. DẠNG 1: PANDOC (WORD NATIVE EQUATION) ---

def convert_md_to_docx_via_pandoc(md_text, temp_dir):
    output_docx_path = os.path.join(temp_dir, "output_native.docx")
    current_cwd = os.getcwd()
    try:
        os.chdir(temp_dir)
        pypandoc.convert_text(
            source=md_text,
            format="markdown",
            to="docx",
            outputfile=output_docx_path,
            extra_args=["--mathjax"],
        )
        with open(output_docx_path, "rb") as f:
            return f.read()
    finally:
        os.chdir(current_cwd)


# --- 4. DẠNG 2: RAW WORD ---

def convert_json_to_docx_raw_bytes(json_data, uploaded_images_map, json_upload_dir=""):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    pdf_info = json_data.get("pdf_info", [])
    for page in pdf_info:
        for block in page.get("para_blocks", []):
            b_type = block.get("type")

            if b_type in ["text", "title"]:
                p = doc.add_paragraph()
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        span_type = span.get("type")
                        content = span.get("content", "")

                        if span_type == "text":
                            if re.match(r"^Bài\s+\d+", content.strip()):
                                run = p.add_run(content)
                                run.bold = True
                            else:
                                p.add_run(content)
                        elif span_type == "inline_equation":
                            run = p.add_run(f" {format_latex_string(content)} ")
                            run.font.name = "Cambria Math"

            elif b_type in ["image", "chart", "figure"]:
                paths_to_check = []
                if block.get("image_path"):
                    paths_to_check.append(block.get("image_path"))
                for sub_b in block.get("blocks", []):
                    if sub_b.get("image_path"):
                        paths_to_check.append(sub_b.get("image_path"))
                    for line in sub_b.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("image_path"):
                                paths_to_check.append(span.get("image_path"))
                
                for img_path_str in paths_to_check:
                    img_stream = get_image_bytes(img_path_str, uploaded_images_map, json_upload_dir)
                    if img_stream:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(img_stream, width=Inches(3.5))

            elif b_type == "table":
                for sub_b in block.get("blocks", []):
                    for line in sub_b.get("lines", []):
                        for span in line.get("spans", []):
                            table_html = span.get("html")
                            if table_html:
                                soup = BeautifulSoup(table_html, "html.parser")
                                rows = soup.find_all("tr")
                                if rows:
                                    doc.add_paragraph()
                                    num_cols = max(len(r.find_all(["td", "th"])) for r in rows)
                                    w_tbl = doc.add_table(rows=len(rows), cols=num_cols)
                                    w_tbl.style = "Table Grid"

                                    for r_idx, tr in enumerate(rows):
                                        for c_idx, cell in enumerate(tr.find_all(["td", "th"])):
                                            if c_idx >= num_cols:
                                                break
                                            cp = w_tbl.cell(r_idx, c_idx).paragraphs[0]

                                            for node in cell.children:
                                                if node.name == "eq":
                                                    r = cp.add_run(f" {format_latex_string(node.get_text())} ")
                                                    r.font.name = "Cambria Math"
                                                elif node.name == "img":
                                                    img_src = node.get("src")
                                                    if img_src:
                                                        img_stream = get_image_bytes(img_src, uploaded_images_map, json_upload_dir)
                                                        if img_stream:
                                                            cp.add_run().add_picture(img_stream, width=Inches(2.5))
                                                elif node.name is None:
                                                    cp.add_run(str(node).strip())

    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io.getvalue()


# --- 5. RENDER PREVIEW TÍCH HỢP NÚT COPY CLIPBOARD ---

def render_preview_with_copy(json_data, uploaded_images_map, json_upload_dir=""):
    preview_inner_html = '<div id="content-to-copy" style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">'
    
    pdf_info = json_data.get("pdf_info", [])
    for page in pdf_info:
        for block in page.get("para_blocks", []):
            b_type = block.get("type")

            if b_type in ["text", "title"]:
                p_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        span_type = span.get("type")
                        content = span.get("content", "")
                        if span_type == "text":
                            if re.match(r"^Bài\s+\d+", content.strip()):
                                p_text += f"<b>{content}</b>"
                            else:
                                p_text += content
                        elif span_type == "inline_equation":
                            clean_c = content.strip().replace("$", "")
                            p_text += f" <i>${clean_c}$</i> "
                if p_text.strip():
                    preview_inner_html += f"<p>{p_text.strip()}</p>"

            elif b_type in ["image", "chart", "figure"]:
                paths_to_check = []
                if block.get("image_path"):
                    paths_to_check.append(block.get("image_path"))
                for sub_b in block.get("blocks", []):
                    if sub_b.get("image_path"):
                        paths_to_check.append(sub_b.get("image_path"))
                    for line in sub_b.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("image_path"):
                                paths_to_check.append(span.get("image_path"))
                
                for img_path_str in paths_to_check:
                    img_stream = get_image_bytes(img_path_str, uploaded_images_map, json_upload_dir)
                    if img_stream:
                        encoded = base64.b64encode(img_stream.getvalue()).decode("utf-8")
                        preview_inner_html += f'<div style="text-align: center; margin: 15px 0;"><img src="data:image/png;base64,{encoded}" style="max-width: 400px; border-radius: 6px;" /></div>'

            elif b_type == "table":
                for sub_b in block.get("blocks", []):
                    for line in sub_b.get("lines", []):
                        for span in line.get("spans", []):
                            table_html = span.get("html")
                            if table_html:
                                soup = BeautifulSoup(table_html, "html.parser")
                                for eq_tag in soup.find_all("eq"):
                                    eq_text = eq_tag.get_text().strip().replace("$", "")
                                    eq_tag.string = f"${eq_text}$"
                                
                                for img_tag in soup.find_all("img"):
                                    img_src = img_tag.get("src")
                                    if img_src:
                                        img_stream = get_image_bytes(img_src, uploaded_images_map, json_upload_dir)
                                        if img_stream:
                                            encoded = base64.b64encode(img_stream.getvalue()).decode("utf-8")
                                            img_tag['src'] = f"data:image/png;base64,{encoded}"
                                            img_tag['width'] = "150"
                                preview_inner_html += f'<div style="margin: 15px 0; overflow-x: auto;">{str(soup)}</div>'

    preview_inner_html += '</div>'

    copier_component = """
    <div style="margin-bottom: 15px;">
        <button onclick="copyContentToClipboard()" style="padding: 10px 20px; background-color: #2b6cb0; color: white; border: none; border-radius: 6px; cursor: pointer; font-weight: bold; font-size: 14px; box-shadow: 0 2px 5px rgba(0,0,0,0.1);">
            📋 Sao chép nội dung Preview (Dán thẳng vào Word)
        </button>
        <span id="copy-status" style="margin-left: 10px; color: #2f855a; font-weight: bold; font-size: 13px; display: none;">✔ Đã sao chép thành công!</span>
    </div>
    
    <div style="background-color: #ffffff; padding: 25px; border-radius: 10px; border: 1px solid #cbd5e0; max-height: 550px; overflow-y: auto;">
        __PREVIEW_INNER_HTML__
    </div>

    <script>
    function copyContentToClipboard() {
        const range = document.createRange();
        range.selectNode(document.getElementById('content-to-copy'));
        window.getSelection().removeAllRanges();
        window.getSelection().addRange(range);
        
        try {
            document.execCommand('copy');
            const status = document.getElementById('copy-status');
            status.style.display = 'inline';
            setTimeout(() => { status.style.display = 'none'; }, 3000);
        } catch (err) {
            alert('Không thể sao chép tự động!');
        }
        window.getSelection().removeAllRanges();
    }
    </script>
    """.replace("__PREVIEW_INNER_HTML__", preview_inner_html)
    
    st.markdown("### 👁️ Xem trước nội dung & Sao chép nhanh")
    components.html(copier_component, height=620, scrolling=False)


# --- 6. GIAO DIỆN STREAMLIT CHÍNH (Đưa phần tải file ra màn hình chính) ---

st.markdown("<h1 style='color: #1a202c;'>🚀 MinerU JSON Converter Pro</h1>", unsafe_allow_html=True)
st.write("Tải lên file JSON và các ảnh đi kèm để chuyển đổi cấu trúc tài liệu sang Word, Markdown hoặc sao chép trực tiếp.")

# Tạo khung tải file ngay ngoài màn hình chính để dễ nhìn, không sợ ẩn Sidebar
st.markdown("---")
col_up1, col_up2 = st.columns(2)
with col_up1:
    uploaded_file = st.file_uploader("📥 Chọn file JSON từ máy tính", type=["json"])
with col_up2:
    uploaded_image_files = st.file_uploader("🖼️ Tải lên thư mục ảnh đi kèm (nếu có)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

uploaded_images_map = {}
if uploaded_image_files:
    for img_file in uploaded_image_files:
        uploaded_images_map[img_file.name] = img_file.getvalue()
    st.success(f"Đã nạp thành công {len(uploaded_image_files)} file ảnh!", icon="✅")

st.markdown("---")

# --- XỬ LÝ KHI CÓ FILE JSON ---
if uploaded_file is not None:
    try:
        json_data = json.load(uploaded_file)
        base_name = uploaded_file.name.rsplit(".", 1)[0]
        json_upload_dir = os.getcwd()
        
        st.success(f"Đang xử lý file thành công: **{uploaded_file.name}**", icon="🚀")

        with tempfile.TemporaryDirectory() as temp_dir:
            with st.spinner("Đang biên dịch cấu trúc tài liệu..."):
                md_content = convert_json_to_markdown(json_data, uploaded_images_map, temp_dir, json_upload_dir)

            st.markdown("### 📥 Các định dạng xuất file")
            col1, col2, col3 = st.columns(3)

            with col1:
                st.markdown("##### 1. Word (Native Math)")
                st.caption("Công thức chuẩn Word Equation")
                with st.spinner("Pandoc đang biên dịch..."):
                    docx_pandoc_bytes = convert_md_to_docx_via_pandoc(md_content, temp_dir)

                st.download_button(
                    label="📥 Tải Word (Native)",
                    data=docx_pandoc_bytes,
                    file_name=f"{base_name}_NativeMath.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                )

            with col2:
                st.markdown("##### 2. Word (Dạng $...$)")
                st.caption("Thích hợp dùng với MathType")
                docx_raw_bytes = convert_json_to_docx_raw_bytes(json_data, uploaded_images_map, json_upload_dir)

                st.download_button(
                    label="📥 Tải Word (Raw Math)",
                    data=docx_raw_bytes,
                    file_name=f"{base_name}_RawMath.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

            with col3:
                st.markdown("##### 3. File Markdown")
                st.caption("Dùng cho Notion / Obsidian")
                st.download_button(
                    label="📥 Tải File Markdown",
                    data=md_content,
                    file_name=f"{base_name}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

            st.divider()

            # Gọi hiển thị xem trước kèm nút Copy
            render_preview_with_copy(json_data, uploaded_images_map, json_upload_dir)

    except Exception as e:
        st.error(f"Đã xảy ra lỗi khi xử lý: {e}")
else:
    st.info("👆 Vui lòng bấm vào nút **Browse files** ở trên để tải lên file **JSON** kết quả từ MinerU.")